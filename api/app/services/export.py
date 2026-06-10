"""Render a persisted memo to PDF or DOCX.

Both render the prose (flagging unverified claims) plus the deterministic data
appendix — the per-fund numbers that ground the memo. Same MemoOut the audit
view consumes.
"""

from __future__ import annotations

import io

from app.schemas.catalog import FundFacts
from app.schemas.memo import MemoOut

# Map common unicode punctuation the LLM emits to latin-1 for PDF core fonts.
_PUNCT = {
    "–": "-", "—": "-", "‘": "'", "’": "'",
    "“": '"', "”": '"', "…": "...", "•": "-",
}


def _latin1(s: str) -> str:
    for a, b in _PUNCT.items():
        s = s.replace(a, b)
    return s.encode("latin-1", "replace").decode("latin-1")


def _get(facts: list, name: str) -> str:
    return next((f.display for f in facts if f.name == name), "—")


def _unverified_count(memo: MemoOut) -> int:
    return sum(1 for s in memo.sections for c in s.claims if not c.verified)


def _status_line(memo: MemoOut) -> str:
    when = memo.created_at.strftime("%Y-%m-%d")
    n = _unverified_count(memo)
    verdict = "all claims grounded" if n == 0 else f"{n} unverified claim(s)"
    return f"Generated {when} · {memo.model} · {verdict}"


def _appendix_row(ff: FundFacts) -> list[str]:
    return [
        str(ff.rank),
        ff.fund_name,
        "Shortlisted" if ff.passed else "Excluded",
        _get(ff.fields, "strategy"),
        _get(ff.fields, "aum_usd"),
        _get(ff.metrics, "annualized_volatility"),
        _get(ff.metrics, "sharpe"),
        _get(ff.metrics, "max_drawdown"),
    ]


_APPENDIX_COLS = ["#", "Fund", "Status", "Strategy", "AUM", "Vol", "Sharpe", "Max DD"]


def memo_to_pdf(memo: MemoOut) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def line(h: float, txt: str, style: str = "", size: int = 11) -> None:
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, h, _latin1(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    line(10, "Investment Committee Memo", "B", 16)
    line(6, _status_line(memo), "", 9)
    pdf.ln(2)

    for section in memo.sections:
        line(8, section.title, "B", 13)
        for claim in section.claims:
            prefix = "" if claim.verified else "[UNVERIFIED] "
            line(6, f"- {prefix}{claim.text}")
        pdf.ln(2)

    line(8, "Data Appendix", "B", 13)
    for ff in memo.appendix:
        cells = _appendix_row(ff)
        line(6, f"{cells[0]}. {cells[1]} ({cells[2]})", "B", 10)
        line(5, "   " + " · ".join(
            f"{c}: {v}" for c, v in zip(_APPENDIX_COLS[3:], cells[3:])
        ), "", 9)

    return bytes(pdf.output())


def memo_to_docx(memo: MemoOut) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Investment Committee Memo", level=0)
    doc.add_paragraph(_status_line(memo))

    for section in memo.sections:
        doc.add_heading(section.title, level=1)
        for claim in section.claims:
            p = doc.add_paragraph(style="List Bullet")
            if not claim.verified:
                run = p.add_run("[UNVERIFIED] ")
                run.bold = True
            p.add_run(claim.text)

    doc.add_heading("Data Appendix", level=1)
    table = doc.add_table(rows=1, cols=len(_APPENDIX_COLS))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(_APPENDIX_COLS):
        table.rows[0].cells[i].text = col
    for ff in memo.appendix:
        cells = table.add_row().cells
        for i, val in enumerate(_appendix_row(ff)):
            cells[i].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
